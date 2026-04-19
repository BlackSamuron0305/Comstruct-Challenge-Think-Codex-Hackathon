"""CSV / XLSX / PDF / Image parsing for the supplier-ingestion pipeline."""
from __future__ import annotations

import io
import logging
import re
from typing import Any

import pandas as pd
from pypdf import PdfReader

MAX_SAMPLE_VALUES = 5
MARKDOWN_DIRECT_CONFIDENCE = 0.66
DEFAULT_PAGES_PER_CHUNK = 20
DEFAULT_PAGE_OVERLAP = 1
MAX_MARKDOWN_CHARS_PER_CHUNK = 32_000
_ALNUM_RE = re.compile(r"[A-Za-z0-9]")
logger = logging.getLogger(__name__)


def parse_tabular(filename: str, content: bytes) -> pd.DataFrame:
    name = filename.lower()
    if name.endswith(".csv"):
        return pd.read_csv(io.BytesIO(content), sep=None, engine="python", dtype=str).fillna("")
    if name.endswith(".tsv"):
        return pd.read_csv(io.BytesIO(content), sep="\t", dtype=str).fillna("")
    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(io.BytesIO(content), dtype=str).fillna("")
    if name.endswith(".ods"):
        return pd.read_excel(io.BytesIO(content), engine="odf", dtype=str).fillna("")
    if name.endswith((".docx", ".doc")):
        return _parse_docx_to_df(content)
    raise ValueError(f"Unsupported file type for tabular ingest: {filename}")


def _parse_docx_to_df(content: bytes) -> pd.DataFrame:
    """Extract tables from a DOCX file; falls back to paragraph text if no tables."""
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(io.BytesIO(content))
    all_rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            all_rows.append([cell.text.strip() for cell in row.cells])

    if all_rows:
        header = all_rows[0]
        data = all_rows[1:] if len(all_rows) > 1 else []
        return pd.DataFrame(data, columns=header).fillna("")

    # Fallback: use paragraph lines as raw text
    lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return pd.DataFrame({"text": lines})


def parse_pdf_to_table(content: bytes) -> pd.DataFrame:
    """Extract tables from PDF using pdfplumber; fall back to pypdf text lines."""
    try:
        import pdfplumber

        all_rows: list[list[str]] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cleaned = [str(cell).strip() if cell else "" for cell in row]
                        if any(cleaned):
                            all_rows.append(cleaned)

        if all_rows:
            header = all_rows[0]
            data = all_rows[1:] if len(all_rows) > 1 else []
            return pd.DataFrame(data, columns=header).fillna("")
    except Exception as e:
        logger.warning("pdfplumber extraction failed (%s), falling back to pypdf", e)

    # Fallback: plain text extraction via pypdf
    reader = PdfReader(io.BytesIO(content))
    rows: list[str] = []
    for page in reader.pages:
        for line in (page.extract_text() or "").splitlines():
            line = line.strip()
            if line:
                rows.append(line)
    return pd.DataFrame({"text": rows})


def parse_pdf_to_markdown_pages(content: bytes) -> list[dict[str, Any]]:
    """Convert a PDF into page-scoped markdown blocks."""
    try:
        import os
        import tempfile
        import pymupdf4llm

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            md_chunks = pymupdf4llm.to_markdown(tmp_path, page_chunks=True)
            pages: list[dict[str, Any]] = []
            if isinstance(md_chunks, list):
                for idx, chunk in enumerate(md_chunks, start=1):
                    if isinstance(chunk, dict):
                        page_no = int(chunk.get("page") or chunk.get("page_number") or chunk.get("number") or idx)
                        text = str(chunk.get("text") or chunk.get("md") or chunk.get("content") or "").strip()
                    else:
                        page_no = idx
                        text = str(chunk).strip()
                    if text:
                        if not text.lstrip().startswith("#"):
                            text = f"## Page {page_no}\n\n{text}"
                        pages.append({"page": page_no, "markdown": text})
            if pages:
                return pages
        finally:
            os.unlink(tmp_path)
    except Exception as e:
        logger.warning("pymupdf4llm page markdown extraction failed (%s), falling back to text", e)

    reader = PdfReader(io.BytesIO(content))
    pages: list[dict[str, Any]] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append({"page": idx, "markdown": f"## Page {idx}\n\n{text}"})
    return pages


def parse_pdf_to_markdown(content: bytes) -> str:
    """Convert PDF to markdown first, with graceful fallback."""
    return "\n\n".join(page["markdown"] for page in parse_pdf_to_markdown_pages(content)).strip()


def build_markdown_chunks(
    page_markdowns: list[dict[str, Any]],
    *,
    pages_per_chunk: int = DEFAULT_PAGES_PER_CHUNK,
    overlap_pages: int = DEFAULT_PAGE_OVERLAP,
    max_chars: int = MAX_MARKDOWN_CHARS_PER_CHUNK,
) -> list[dict[str, Any]]:
    """Build page-aware markdown chunks with overlap and token-aware sizing."""
    if not page_markdowns:
        return []

    pages_per_chunk = max(1, int(pages_per_chunk))
    overlap_pages = max(0, min(int(overlap_pages), pages_per_chunk - 1))
    max_chars = max(2_000, int(max_chars))

    chunks: list[dict[str, Any]] = []
    index = 0
    total_pages = len(page_markdowns)

    while index < total_pages:
        selected: list[dict[str, Any]] = []
        char_count = 0
        cursor = index

        while cursor < total_pages and len(selected) < pages_per_chunk:
            page_info = page_markdowns[cursor]
            markdown = str(page_info.get("markdown") or "").strip()
            if not markdown:
                cursor += 1
                continue
            if selected and char_count + len(markdown) > max_chars:
                break
            selected.append({
                "page": int(page_info.get("page") or cursor + 1),
                "markdown": markdown,
            })
            char_count += len(markdown)
            cursor += 1

        if not selected:
            page_info = page_markdowns[index]
            selected = [{
                "page": int(page_info.get("page") or index + 1),
                "markdown": str(page_info.get("markdown") or "").strip(),
            }]
            cursor = index + 1

        chunks.append({
            "start_page": selected[0]["page"],
            "end_page": selected[-1]["page"],
            "page_numbers": [entry["page"] for entry in selected],
            "markdown": "\n\n".join(entry["markdown"] for entry in selected if entry["markdown"]),
            "char_count": char_count,
        })

        if cursor >= total_pages:
            break

        next_index = cursor - overlap_pages
        if next_index <= index:
            next_index = index + 1
        index = next_index

    return chunks


def extract_catalog_from_markdown(markdown: str) -> list[dict]:
    """Try deterministic extraction from markdown tables before LLM fallback."""
    items: list[dict] = []
    lines = [ln.strip() for ln in markdown.splitlines() if ln.strip()]
    for line in lines:
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 2:
            continue
        joined = " ".join(cells).lower()
        if "sku" in joined and "price" in joined:
            continue
        if set("".join(cells)) <= {"-", ":"}:
            continue
        sku = cells[0] if _ALNUM_RE.search(cells[0]) else ""
        name = cells[1] if len(cells) > 1 else ""
        if not name:
            continue
        price_match = re.search(r"(\d+[.,]?\d*)", " ".join(cells[2:])) if len(cells) > 2 else None
        unit_price = float(price_match.group(1).replace(",", ".")) if price_match else None
        items.append({
            "sku": sku or None,
            "name": name,
            "unit_price": unit_price,
            "currency": "CHF",
            "confidence": MARKDOWN_DIRECT_CONFIDENCE,
            "extraction_mode": "markdown_direct",
        })
    return items


def parse_image_to_text(content: bytes) -> str:
    """Extract text from an image using Tesseract OCR."""
    from PIL import Image
    import pytesseract

    image = Image.open(io.BytesIO(content))
    # Use German + English + French for Swiss construction docs
    text = pytesseract.image_to_string(image, lang="deu+eng+fra")
    return text.strip()


def parse_image_to_table(content: bytes) -> pd.DataFrame:
    """Extract text from image via OCR, return as single-column DataFrame for LLM processing."""
    text = parse_image_to_text(content)
    if not text:
        return pd.DataFrame()
    rows = [line.strip() for line in text.splitlines() if line.strip()]
    return pd.DataFrame({"text": rows})


def column_samples(df: pd.DataFrame) -> list[dict[str, Any]]:
    """Returns [{"name": col, "samples": [..]}] suitable for the column_mapper prompt."""
    out: list[dict[str, Any]] = []
    for col in df.columns:
        samples = (
            df[col].astype(str).str.strip().replace("", pd.NA).dropna().unique().tolist()
        )
        out.append({"name": str(col), "samples": samples[:MAX_SAMPLE_VALUES]})
    return out


def apply_mapping(df: pd.DataFrame, mappings: list[dict]) -> list[dict]:
    """Apply the selected database-field mapping and preserve extra detail in special_info."""
    rows: list[dict] = []

    for _, df_row in df.iterrows():
        mapped_row: dict[str, Any] = {}
        special_info: dict[str, Any] = {}

        for mapping in mappings:
            source_column = str(mapping.get("source_column") or "")
            target_field = mapping.get("target_field")
            if not source_column or not target_field or source_column not in df.columns:
                continue

            raw_value = df_row.get(source_column)
            if pd.isna(raw_value):
                continue

            value = str(raw_value).strip()
            if value == "":
                continue

            if target_field == "special_info":
                special_info[source_column] = value
                continue

            if target_field not in mapped_row or mapped_row[target_field] in (None, ""):
                mapped_row[target_field] = value
            else:
                special_info[source_column] = value

        if special_info:
            existing_special = mapped_row.get("special_info")
            if isinstance(existing_special, dict):
                mapped_row["special_info"] = {**existing_special, **special_info}
            elif existing_special not in (None, ""):
                mapped_row["special_info"] = {"note": existing_special, **special_info}
            else:
                mapped_row["special_info"] = special_info

        rows.append(mapped_row)

    return rows
